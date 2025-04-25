#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import docx
from docx.document import Document
from docx.oxml.text.run import CT_R
from docx.oxml.text.paragraph import CT_P
from zipfile import ZipFile
import win32com.client
import pythoncom

def check_word_file(file_path, file_name, enable_value_search=False, search_values=None):
    """
    Проверка Word файла
    
    Args:
        file_path (str): Путь к файлу
        file_name (str): Имя файла
        enable_value_search (bool): Флаг включения поиска заданных значений
        search_values (list): Список значений для поиска
    """
    try:
        issues = []
        found_values = []
        
        # Для DOCX используем стандартный подход
        if file_path.lower().endswith('.docx'):
            # Пытаемся использовать быструю проверку файла без полной загрузки для DOCX
            try:
                # Предварительная проверка на комментарии по содержимому XML
                from zipfile import ZipFile
                with ZipFile(file_path) as docx_zip:
                    # Проверяем наличие файла comments.xml
                    if 'word/comments.xml' in docx_zip.namelist():
                        issues.append("комментарии")
                    
                    # Проверяем наличие выделений определенными цветами в document.xml
                    if not issues and 'word/document.xml' in docx_zip.namelist():
                        # Считываем первые 100 КБ файла для базовой проверки
                        document_content = docx_zip.read('word/document.xml', 100 * 1024)  # Первые 100 KB
                        document_content_str = document_content.decode('utf-8', errors='ignore')
                        
                        # Ищем теги highlight с определенными значениями цвета
                        if 'w:highlight w:val="yellow"' in document_content_str:
                            issues.append("желтые выделения")
                        elif 'w:highlight w:val="red"' in document_content_str:
                            issues.append("красные выделения")
                        elif 'w:highlight w:val="green"' in document_content_str:
                            issues.append("зеленые выделения")
                        elif 'w:highlight w:val="blue"' in document_content_str:
                            issues.append("синие выделения")
                    
                    # Если включен поиск значений, проверяем document.xml
                    if enable_value_search and search_values and 'word/document.xml' in docx_zip.namelist():
                        # Считываем содержимое документа для поиска
                        document_content = docx_zip.read('word/document.xml')
                        document_content_str = document_content.decode('utf-8', errors='ignore')
                        
                        # Ищем заданные значения
                        for value in search_values:
                            if value in document_content_str and value not in found_values:
                                found_values.append(value)
            except Exception:
                # Если быстрая проверка не удалась, переходим к обычной проверке
                pass
            
            # Если мы все еще не обнаружили проблем, используем стандартный подход для DOCX
            if not issues or "желтые выделения" not in issues and "красные выделения" not in issues and "зеленые выделения" not in issues and "синие выделения" not in issues:
                try:
                    doc = docx.Document(file_path)
                    
                    # Улучшенная проверка на комментарии
                    if not "комментарии" in issues:
                        try:
                            # Проверка через xpath
                            comments = doc._part._element.body.xpath('//w:commentRangeStart') + doc._part._element.body.xpath('//w:commentRangeEnd')
                            if comments and len(comments) > 0:
                                issues.append("комментарии")
                            # Альтернативная проверка через core.xml
                            elif hasattr(doc, '_part') and hasattr(doc._part, 'comments_part') and doc._part.comments_part:
                                issues.append("комментарии")
                        except:
                            # Если не смогли проверить первым способом, пробуем альтернативный
                            try:
                                # Проверка в document.xml
                                if "commentReference" in doc._element.xml or "commentRangeStart" in doc._element.xml:
                                    issues.append("комментарии")
                            except:
                                pass
                    
                    # Проверка на выделения определенными цветами
                    if "желтые выделения" not in issues and "красные выделения" not in issues and "зеленые выделения" not in issues and "синие выделения" not in issues:
                        # Проверяем только первые 100 параграфов для скорости
                        max_paragraphs = min(100, len(doc.paragraphs))
                        
                        for idx, paragraph in enumerate(doc.paragraphs):
                            if idx >= max_paragraphs:
                                break
                                
                            for run in paragraph.runs:
                                if run.element.rPr is not None:
                                    # Проверяем через xpath с фильтрацией по определенным цветам
                                    try:
                                        # Проверяем наличие выделения и определяем его цвет
                                        highlight_nodes = run.element.rPr.xpath('.//w:highlight')
                                        
                                        if highlight_nodes:
                                            for highlight in highlight_nodes:
                                                # Получаем значение атрибута val
                                                val = highlight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                                
                                                if val == 'yellow':
                                                    issues.append("желтые выделения")
                                                    break
                                                elif val == 'red':
                                                    issues.append("красные выделения")
                                                    break
                                                elif val == 'green':
                                                    issues.append("зеленые выделения")
                                                    break
                                                elif val == 'blue':
                                                    issues.append("синие выделения")
                                                    break
                                    except:
                                        # Если xpath не сработал, пробуем другой подход
                                        try:
                                            if hasattr(run.font, 'highlight_color'):
                                                color = run.font.highlight_color
                                                if color == 'yellow':
                                                    issues.append("желтые выделения")
                                                    break
                                                elif color == 'red':
                                                    issues.append("красные выделения")
                                                    break
                                                elif color == 'green':
                                                    issues.append("зеленые выделения")
                                                    break
                                                elif color == 'blue':
                                                    issues.append("синие выделения")
                                                    break
                                        except:
                                            pass
                                
                                # Если уже нашли выделения, завершаем проверку
                                if any(item in issues for item in ["желтые выделения", "красные выделения", "зеленые выделения", "синие выделения"]):
                                    break
                            
                            # Если уже нашли выделения, завершаем проверку
                            if any(item in issues for item in ["желтые выделения", "красные выделения", "зеленые выделения", "синие выделения"]):
                                break
                    
                    # Поиск заданных пользователем значений
                    if enable_value_search and search_values and not found_values:
                        max_paragraphs = min(500, len(doc.paragraphs))  # Увеличиваем количество параграфов для поиска
                        
                        for idx, paragraph in enumerate(doc.paragraphs):
                            if idx >= max_paragraphs:
                                break
                                
                            # Проверяем текст параграфа на наличие искомых значений
                            if paragraph.text:
                                for value in search_values:
                                    if value in paragraph.text and value not in found_values:
                                        found_values.append(value)
                            
                            # Если нашли все значения, прекращаем поиск
                            if len(found_values) == len(search_values):
                                break
                                
                except Exception as e:
                    # Если не удалось открыть файл через docx.Document, используем win32com
                    return check_word_file_with_win32com(file_path, file_name, enable_value_search, search_values)
                    
        # Для .docm используем упрощенный подход с Win32COM
        elif file_path.lower().endswith('.docm'):
            # Упрощенный вариант проверки для DOCM
            try:
                # Для DOCM сразу пробуем использовать упрощенную проверку через ZipFile
                from zipfile import ZipFile
                try:
                    with ZipFile(file_path) as docm_zip:
                        # Проверяем наличие файла comments.xml
                        if 'word/comments.xml' in docm_zip.namelist():
                            issues.append("комментарии")
                        
                        # Проверяем наличие выделений определенными цветами в document.xml
                        if 'word/document.xml' in docm_zip.namelist():
                            # Считываем первые 100 КБ файла для базовой проверки
                            document_content = docm_zip.read('word/document.xml', 100 * 1024)
                            document_content_str = document_content.decode('utf-8', errors='ignore')
                            
                            # Ищем теги highlight с определенными значениями цвета
                            if 'w:highlight w:val="yellow"' in document_content_str:
                                issues.append("желтые выделения")
                            elif 'w:highlight w:val="red"' in document_content_str:
                                issues.append("красные выделения")
                            elif 'w:highlight w:val="green"' in document_content_str:
                                issues.append("зеленые выделения")
                            elif 'w:highlight w:val="blue"' in document_content_str:
                                issues.append("синие выделения")
                            
                        # Если включен поиск значений, проверяем document.xml
                        if enable_value_search and search_values and 'word/document.xml' in docm_zip.namelist():
                            # Считываем содержимое документа для поиска
                            document_content = docm_zip.read('word/document.xml')
                            document_content_str = document_content.decode('utf-8', errors='ignore')
                            
                            # Ищем заданные значения
                            for value in search_values:
                                if value in document_content_str and value not in found_values:
                                    found_values.append(value)
                except Exception:
                    # Если ZipFile не сработал, используем Win32COM с упрощенной проверкой
                    return check_word_file_with_win32com(file_path, file_name, enable_value_search, search_values)
            except Exception:
                # При любой ошибке используем Win32COM с упрощенной проверкой
                return check_word_file_with_win32com(file_path, file_name, enable_value_search, search_values)
                
        # Для .doc всегда используем win32com
        elif file_path.lower().endswith('.doc'):
            return check_word_file_with_win32com(file_path, file_name, enable_value_search, search_values)
            
        # Если найдены указанные значения, добавляем их в проблемы
        if found_values:
            issues.append(f"найдены заданные значения: {', '.join(found_values)}")
            
        # Формирование результата
        if issues:
            result = "Не пройден"
            comment = "Найдено: " + ", ".join(set(issues))
        else:
            result = "Пройден"
            comment = "Проблем не обнаружено"
        
        return {
            'file_name': file_name,
            'file_type': "Word",
            'file_path': file_path,
            'result': result,
            'comment': comment
        }
        
    except Exception as e:
        return {
            'file_name': file_name,
            'file_type': "Word",
            'file_path': file_path,
            'result': "Ошибка",
            'comment': f"Ошибка проверки: {str(e)}"
        }

def check_word_file_with_win32com(file_path, file_name, enable_value_search=False, search_values=None):
    """
    Упрощенная проверка Word файла с использованием win32com
    
    Args:
        file_path (str): Путь к файлу
        file_name (str): Имя файла
        enable_value_search (bool): Флаг включения поиска заданных значений
        search_values (list): Список значений для поиска
    """
    try:
        issues = []
        found_values = []
        pythoncom.CoInitialize()
        word_app = None
        doc = None
        
        try:
            # Создаем экземпляр Word с минимальными настройками
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False  # Отключаем диалоги с предупреждениями
            word_app.AutomationSecurity = 3  # msoAutomationSecurityForceDisable - отключаем макросы
            
            # Пытаемся открыть документ с минимальным доступом
            try:
                doc = word_app.Documents.Open(
                    file_path, 
                    ReadOnly=True,              # Только для чтения
                    AddToRecentFiles=False,     # Не добавлять в список недавних
                    Visible=False,              # Невидимый режим
                    OpenAndRepair=False,        # Не пытаться восстанавливать
                    DoNotLoadVbaAndOpenDocm=True # Не загружать VBA для DOCM
                )
                
                # УПРОЩЕННАЯ ПРОВЕРКА: только комментарии и цвета выделений
                
                # 1. Проверка на комментарии - простой подход
                try:
                    if doc.Comments.Count > 0:
                        issues.append("комментарии")
                except Exception:
                    pass  # Пропускаем, если не можем проверить комментарии
                
                # 2. Проверка на цветные выделения - ограничиваем область проверки
                try:
                    # Константы цветов из Word
                    wdYellow = 7  # желтый
                    wdRed = 6     # красный
                    wdGreen = 4   # зеленый
                    wdBlue = 9    # синий
                    
                    # Проверяем только несколько первых абзацев
                    max_paragraphs = 10
                    paragraphs_count = min(max_paragraphs, doc.Paragraphs.Count)
                    
                    for i in range(1, paragraphs_count + 1):
                        try:
                            paragraph = doc.Paragraphs(i)
                            
                            # Проверяем выделение для всего абзаца
                            if paragraph.Range.HighlightColorIndex == wdYellow:
                                issues.append("желтые выделения")
                                break
                            elif paragraph.Range.HighlightColorIndex == wdRed:
                                issues.append("красные выделения")
                                break
                            elif paragraph.Range.HighlightColorIndex == wdGreen:
                                issues.append("зеленые выделения")
                                break
                            elif paragraph.Range.HighlightColorIndex == wdBlue:
                                issues.append("синие выделения")
                                break
                        except Exception:
                            continue  # Пропускаем проблемные абзацы
                except Exception:
                    pass  # Пропускаем всю проверку выделений, если есть проблемы
                
                # 3. Поиск заданных пользователем значений
                if enable_value_search and search_values:
                    try:
                        # Определяем максимальное количество параграфов для поиска
                        max_search_paragraphs = min(100, doc.Paragraphs.Count)
                        
                        for i in range(1, max_search_paragraphs + 1):
                            try:
                                paragraph = doc.Paragraphs(i)
                                if paragraph.Range.Text:
                                    paragraph_text = paragraph.Range.Text
                                    
                                    # Ищем заданные значения в тексте абзаца
                                    for value in search_values:
                                        if value in paragraph_text and value not in found_values:
                                            found_values.append(value)
                                
                                # Если нашли все значения, прекращаем поиск
                                if len(found_values) == len(search_values):
                                    break
                            except Exception:
                                continue  # Пропускаем проблемные абзацы
                    except Exception:
                        # Если не удалось провести поиск по абзацам, пробуем альтернативный подход
                        try:
                            # Проверяем весь текст документа сразу
                            doc_text = doc.Content.Text
                            for value in search_values:
                                if value in doc_text and value not in found_values:
                                    found_values.append(value)
                        except Exception:
                            pass  # Игнорируем ошибки при альтернативном подходе
                
            except pythoncom.com_error as com_err:
                # Упрощаем обработку ошибок - единый формат без деталей
                return {
                    'file_name': file_name,
                    'file_type': "Word",
                    'file_path': file_path,
                    'result': "Предупреждение",
                    'comment': "Файл содержит защищенные элементы (возможно, макросы). Базовая проверка невозможна."
                }
                    
        finally:
            # Гарантированное закрытие документа и приложения Word
            try:
                if doc is not None:
                    doc.Close(SaveChanges=False)
            except:
                pass
                
            try:
                if word_app is not None:
                    word_app.Quit()
            except:
                pass
        
        # Если найдены указанные значения, добавляем их в проблемы
        if found_values:
            issues.append(f"найдены заданные значения: {', '.join(found_values)}")
                    
        # Формирование результата
        if issues:
            result = "Не пройден"
            comment = "Найдено: " + ", ".join(set(issues))
        else:
            result = "Пройден"
            comment = "Проблем не обнаружено"
        
        return {
            'file_name': file_name,
            'file_type': "Word",
            'file_path': file_path,
            'result': result,
            'comment': comment
        }
        
    except Exception as e:
        # Упрощенный общий обработчик ошибок
        return {
            'file_name': file_name,
            'file_type': "Word",
            'file_path': file_path,
            'result': "Пропущен",
            'comment': "Файл содержит сложную структуру. Проверка пропущена."
        }